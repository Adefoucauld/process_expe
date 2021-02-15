# -*- coding: utf-8 -*-
"""
Created on Thu Jan 14 09:53:17 2021

@author: Utilisateur
"""
from datetime import datetime # pour importer la date 
import csv 
import win32api
import shutil

from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_LINE_SPACING
import pandas as pd

#on crée une function que l'on va appeler 
def gene_dico():
    '''Importer liste de shop et constituion des dico'''
    def conversion(L):      # permet le traitement de fichier csv comme une liste
        T=[]
        for i in L:
            T.append(int(i.replace(',','.')))
        return(T)
        
    num_shop =[]
    address_shop =[]
    zip_shop =[]
    town_shop =[]
    
    #on récupère le fichier excel des adresses des shops vdb où on a ajouté les shops fnac livrés par vdb
    with open('C:/Users/Utilisateur/Documents/FNAC_VDB/Projet_expe/Copie de WEB_LIST_SHOP.csv', 'rt') as csvfile:  # on importe les adresses des shops
        spamreader = csv.reader(csvfile,delimiter = ';')
        for row in spamreader:
            num_shop.append(row[1])
            address_shop.append(row[6])
            zip_shop.append(row[7])
            town_shop.append(row[8])
              
    num_shop=conversion(num_shop) 
    
    dico_add = {}                #on crée nos dico d'adresses
    dico_zip = {}
    for i in range(len(num_shop)):
        addresse = address_shop[i] +'\n'+ zip_shop[i] +' '+ town_shop[i]
        dico_add[num_shop[i]] = addresse
        dico_zip[num_shop[i]] = zip_shop[i]
        # print(dico_add[num_shop[i]])
    dico_add['SPL'] = "Fnac Vanden Borre"+'\n'+"Slesbroekstraat 101" +'\n'+ "B-1600 Sint-Pieters-Leeuw"
    dico_zip['SPL'] = "B-1600"
    
    return(num_shop,address_shop,zip_shop,town_shop,dico_add,dico_zip)


''' Génération du CMR pour un shop '''



def gene_cmr(R,pal,shop_pal):
    
    num_shop,address_shop,zip_shop,town_shop,dico_add,dico_zip = gene_dico()
    
    #On récup la date et l'heure en focntion du format à afficher
    now = datetime.now()
    date_time = now.strftime("%d/%m/%Y     %H:%M")
    date_only = now.strftime("%d/%m/%Y     ..... h .....")
    datecode = now.strftime("%Y%m%d")
    date_CMR = now.strftime("_%d%m%Y")
    
    cmr_autor = True
    
    keys = list(R.keys())
    keys.remove('tournee')
    keys.remove('dacem')
    
    #on crée un dico pour actualiser le fichier de palette après coup
    reste_quai = dict()
    for num in keys:
        '''on regarde les restes à quais'''
        
        reste_quai[num]=[]
        for j in range(len(pal[num])):
            if pal[num][j] not in R[num]:
                reste_quai[num].append(pal[num][j])
        pal[num]=reste_quai[num]   #on actualise pal avec les palettes qui restent
        
        
        #on génère le fichier cmr
        if len(R[num])>0:
            document = Document()
            sections = document.sections
            
            for section in sections:     #on passe en marges fines
                section.top_margin = Cm(0.5)
                section.bottom_margin = Cm(0.5)
                section.left_margin = Cm(1)
                section.right_margin = Cm(1)
            
            document.add_heading('FEUILLE DE ROUTE/VRACHTBRIEF', 1)
            chapo = document.add_table(rows = 1, cols = 2)
            chapo.cell(0,0).text = dico_add['SPL']+ '\n02/334.00.00\nTVA/BTW:BE412723419'
            chapo.cell(0,1).text = datecode +'S'+str(num)+'\nEdition:' + date_time + '\nTournée/Ronde : '+str(R['tournee'][0])+' \n'  +'\nShop: Shopping('+str(num)+')'
            
            p = document.add_paragraph()
            p.paragraph_format.line_spacing = 0.7
            p.add_run("Ce transport est soumis, nonobstant toute clause contraire, à la Convention CMR\nDit vervoer is , ongeacht enig tegenstrijdig beding, onderworpen aa het CMR- Verdrag").bold = True
            
            
            def make_rows_bold(*rows):  #on mets des lignes en gras
                for row in rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                
            table1 = document.add_table(rows =6, cols = 2)
            
            table1.cell(0,0).text = '1.Expéditeur/Afzender' 
            table1.cell(0,1).text = '2.Destinataire/Geadresseerde' 
            table1.cell(1,0).text = dico_add['SPL']
            table1.cell(1,1).text = "Vanden Borre Shopping(" + str(num) + ')\n' + dico_add[num]
            table1.cell(2,0).text = "3.Transporteur/Vervoederder (cachet/stample)"
            table1.cell(2,1).text =  "4a.Prise en charge de la marchandise / Inontvangsteneming van de goedren"
            table1.cell(3,1).text =  dico_zip['SPL'] + '  '+ date_time 
            table1.cell(4,1).text = "4b.Livraison de la marchandise/ Aflevering van de goederen" 
            table1.cell(5,1).text =  dico_zip[num] + '  '+ date_only 
            table1.cell(5,0).text = "5.Plomb/Lood Nr ..........................................................................."
            
            make_rows_bold(table1.rows[0],table1.rows[2],table1.rows[4])
            
            p = document.add_paragraph()
            p.paragraph_format.line_spacing = 0.7
            p.add_run("6.Marchandises transportées/ Vervoerde goederen").bold = True
            
            table = document.add_table(rows=7, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Palette/Palet'
            hdr_cells[1].text = 'Nombre/Aantal'
            hdr_cells[2].text = 'Contenu/Inhoud'
            table.cell(1,0).text = 'DACEM'
            table.cell(2,0).text = 'COMBI'
            table.cell(3,0).text = 'Palette/Palet'
            table.cell(5,0).text  = 'Total Palette/Palet'
            table.cell(5,1).text  = 'Total Colis/Collié'
            table.cell(5,2).text  = 'Total Poids brut/Bruttogewicht'
            table.cell(6,2).text  = 'Max 7,5 T'
            table.cell(3,1).text  = 'Nombre/Aantal'
            table.cell(3,2).text  = 'Contenu/Inhoud'
            table.cell(5,1).text  = 'Total Colis/Collié'
            
            
            table.cell(1,1).text  = str(R['dacem'][num]) #str(R[idx][2])
            table.cell(2,1).text  = str(1)
            table.cell(1,2).text  = 'Accessoires'
            table.cell(2,2).text  = 'SAV'
                  
            
            nb_pal =''
            for i in range(len(R[num])):
                nb_pal += str(R[num][i])+'\n' 
             
                           
            table.cell(4,0).text = nb_pal
            table.cell(4,2).text = 'Electro-Ménager/Elektro-huishoud'
            table.cell(6,0).text = str(len(R[num]))
            
            
            make_rows_bold(table.rows[0],table.rows[3],table.rows[5])
            
            p = document.add_paragraph()
            p.paragraph_format.line_spacing = 0.7
            p.add_run('7.Frais afférent au transport/ Vervoerskosten\n8.Formalités de douane et autres/Douanes en andere formaliteiten').bold = True
            
            tablo = document.add_table(rows=2, cols=3)
            tit_cells = tablo.rows[0].cells
            tit_cells[0].text = 'CHAUFFEUR'
            tit_cells[1].text = 'REMARQUES/OPMERKINGEN'
            tit_cells[2].text = 'SHOP'
            tablo.cell(1,0).text = 'Nom:...........................................\n' + 'Sign:...........................................'
            tablo.cell(1,1).text = '......................................................................\n'+'......................................................................\n'+'......................................................................\n'
            tablo.cell(1,2).text = 'Nom:...........................................\n' + 'Sign:............................................'
            
            document.save('C:/Users/Utilisateur/Documents/FNAC_VDB/Projet_expe/CMR_expé_shop_'+str(num)+date_CMR+'.docx')
            
        
            #on copie le cmr dans un dossier tampon pour l'impression
            filePath = shutil.copy('CMR_expé_shop_'+str(num)+date_CMR+'.docx','C:/Users/Utilisateur/Documents/FNAC_VDB/Projet_expe/cmr_print/')
            
            
